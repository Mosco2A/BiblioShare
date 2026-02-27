import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import lxml.etree as etree

BROWN=RGBColor(139,111,78)
WHITE=RGBColor(255,255,255)
INLINERED=RGBColor(140,80,40)
CODEGRAY=RGBColor(40,40,40)
BROWN_HEX='8B6F4E'
GRAY_HEX='F5F5F5'
LTGRAY_HEX='EEEEEE'


def set_cell_shading(cell,fill_hex):
    tc=cell._tc
    tcPr=tc.get_or_add_tcPr()
    xml='<w:shd '+nsdecls('w')+' w:val="clear" w:color="auto" w:fill="'+fill_hex+'"'+'/>'
    tcPr.append(parse_xml(xml))


def set_para_shading(para,fill_hex):
    pPr=para._p.get_or_add_pPr()
    xml='<w:shd '+nsdecls('w')+' w:val="clear" w:color="auto" w:fill="'+fill_hex+'"'+'/>'
    pPr.append(parse_xml(xml))

def add_page_numbers(doc):
    section=doc.sections[0]
    footer=section.footer
    footer.is_linked_to_previous=False
    para=footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment=WD_ALIGN_PARAGRAPH.CENTER
    para.clear()
    r1=para.add_run()
    fc1=etree.SubElement(r1._r,qn("w:fldChar"))
    fc1.set(qn("w:fldCharType"),"begin")
    r2=para.add_run()
    it=etree.SubElement(r2._r,qn("w:instrText"))
    it.set("{http://www.w3.org/XML/1998/namespace}space","preserve")
    it.text=" PAGE "
    r3=para.add_run()
    fc2=etree.SubElement(r3._r,qn("w:fldChar"))
    fc2.set(qn("w:fldCharType"),"end")
    for r in para.runs:
        r.font.name="Calibri"
        r.font.size=Pt(10)
        r.font.color.rgb=RGBColor(128,128,128)

def style_heading(para,level):
    for run in para.runs:
        run.font.color.rgb=BROWN
        run.font.name="Calibri"
        run.bold=True
        if level==1: run.font.size=Pt(18)
        elif level==2: run.font.size=Pt(14)
        else: run.font.size=Pt(12)

def apply_inline(para,text):
    BT=chr(96)
    PAT=BT+chr(40)+chr(91)+chr(94)+BT+chr(93)+chr(43)+chr(41)+BT
    parts=re.split(PAT,text)
    for idx,part in enumerate(parts):
        if not part: continue
        if idx%2==1:
            r=para.add_run(part)
            r.font.name="Courier New"
            r.font.size=Pt(10)
            r.font.color.rgb=INLINERED
        else:
            r=para.add_run(part)
            r.font.name="Calibri"
            r.font.size=Pt(11)

def add_title_page(doc):
    for _ in range(6):
        p=doc.add_paragraph();p.paragraph_format.space_after=Pt(0)
    tp=doc.add_paragraph();tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=tp.add_run("BiblioShare")
    r.font.name="Calibri";r.font.size=Pt(36);r.font.bold=True;r.font.color.rgb=BROWN
    sp=doc.add_paragraph();sp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r2=sp.add_run("Documentation Technique")
    r2.font.name="Calibri";r2.font.size=Pt(22);r2.font.color.rgb=RGBColor(80,80,80)
    sep=doc.add_paragraph();sep.alignment=WD_ALIGN_PARAGRAPH.CENTER
    sr=sep.add_run(chr(8212)*40)
    sr.font.color.rgb=BROWN;sr.font.size=Pt(14)
    for _ in range(3):
        p=doc.add_paragraph();p.paragraph_format.space_after=Pt(0)
    meta=[
        ("Plateforme","iOS + Android (Flutter)"),
        ("Backend","Firebase Auth + Supabase PostgreSQL"),
        ("Version","1.0.0"),
        ("Date","2026-02-20"),
    ]
    for label,value in meta:
        m=doc.add_paragraph();m.alignment=WD_ALIGN_PARAGRAPH.CENTER
        rl=m.add_run(label+": ")
        rl.font.name="Calibri";rl.font.size=Pt(12);rl.font.bold=True;rl.font.color.rgb=BROWN
        rv=m.add_run(value)
        rv.font.name="Calibri";rv.font.size=Pt(12);rv.font.color.rgb=RGBColor(60,60,60)
    doc.add_page_break()

def flush_table(doc,table_rows_buf):
    if not table_rows_buf: return []
    sep_pat=r"^[|][-| :]+[|]$"
    data_rows=[r for r in table_rows_buf if not re.match(sep_pat,r.strip())]
    if not data_rows: return []
    parsed=[]
    for row in data_rows:
        cells=[c.strip() for c in row.strip().strip("|").split("|")]
        parsed.append(cells)
    if not parsed: return []
    max_cols=max(len(r) for r in parsed)
    tbl=doc.add_table(rows=len(parsed),cols=max_cols)
    tbl.style="Table Grid"
    tbl.alignment=WD_TABLE_ALIGNMENT.LEFT
    for ri,row_data in enumerate(parsed):
        row_obj=tbl.rows[ri]
        for ci in range(max_cols):
            cell=row_obj.cells[ci]
            ct=row_data[ci] if ci<len(row_data) else ""
            BT=chr(96)
            PAT=BT+chr(40)+chr(91)+chr(94)+BT+chr(93)+chr(43)+chr(41)+BT
            plain=re.sub(PAT,r"\g<1>",ct)
            p2=cell.paragraphs[0];p2.clear()
            if ri==0:
                set_cell_shading(cell,BROWN_HEX)
                run=p2.add_run(plain)
                run.font.bold=True;run.font.color.rgb=WHITE
                run.font.name="Calibri";run.font.size=Pt(10)
            else:
                if ri%2==0: set_cell_shading(cell,LTGRAY_HEX)
                run=p2.add_run(plain)
                run.font.name="Calibri";run.font.size=Pt(10)
    spacer=doc.add_paragraph()
    spacer.paragraph_format.space_after=Pt(4)
    return []

def flush_code(doc,code_lines):
    if not code_lines: return
    for line in code_lines:
        p=doc.add_paragraph()
        p.paragraph_format.space_before=Pt(0)
        p.paragraph_format.space_after=Pt(0)
        p.paragraph_format.left_indent=Cm(0.5)
        set_para_shading(p,GRAY_HEX)
        content=line.rstrip(chr(10)) or chr(32)
        run=p.add_run(content)
        run.font.name="Courier New"
        run.font.size=Pt(9)
        run.font.color.rgb=CODEGRAY
    gap=doc.add_paragraph()
    gap.paragraph_format.space_before=Pt(2)
    gap.paragraph_format.space_after=Pt(6)

def parse_and_build(md_path,docx_path):
    with open(md_path,encoding="utf-8") as f:
        lines=f.readlines()
    doc=Document()
    doc.styles["Normal"].font.name="Calibri"
    doc.styles["Normal"].font.size=Pt(11)
    section=doc.sections[0]
    mg=Cm(2.5)
    section.top_margin=mg
    section.bottom_margin=mg
    section.left_margin=mg
    section.right_margin=mg
    add_page_numbers(doc)
    add_title_page(doc)
    i=0
    in_code=False
    code_lang=""
    code_lines=[]
    table_buf=[]
    while i<len(lines):
        line=lines[i]
        stripped=line.rstrip(chr(10))
        FENCE=chr(96)*3
        if stripped.startswith(FENCE):
            if not in_code:
                if table_buf: table_buf=flush_table(doc,table_buf)
                in_code=True
                code_lang=stripped[3:].strip()
                code_lines=[]
            else:
                in_code=False
                flush_code(doc,code_lines)
                code_lines=[];code_lang=chr(34)*0
            i+=1;continue
        if in_code:
            code_lines.append(stripped)
            i+=1;continue
        if re.match(r"^---+$",stripped):
            if table_buf: table_buf=flush_table(doc,table_buf)
            i+=1;continue
        if stripped.startswith("|"):
            table_buf.append(stripped)
            i+=1;continue
        else:
            if table_buf: table_buf=flush_table(doc,table_buf)
        hm=re.match(r"^(#{1,4})\s+(.*)",stripped)
        if hm:
            level=len(hm.group(1))
            text=hm.group(2).strip()
            if level==1: i+=1;continue
            word_level=level-1
            BT=chr(96)
            PAT=BT+chr(40)+chr(91)+chr(94)+BT+chr(93)+chr(43)+chr(41)+BT
            plain=re.sub(PAT,r"\g<1>",text)
            p=doc.add_heading(plain,level=word_level)
            p.paragraph_format.space_before=Pt(14 if word_level==1 else 8)
            p.paragraph_format.space_after=Pt(6)
            style_heading(p,word_level)
            i+=1;continue
        bm=re.match(r"^(\s*)[-*+]\s+(.*)",stripped)
        if bm:
            indent=len(bm.group(1))
            text=bm.group(2)
            sname="List Bullet" if indent==0 else "List Bullet 2"
            p=doc.add_paragraph(style=sname)
            p.paragraph_format.space_before=Pt(0)
            p.paragraph_format.space_after=Pt(2)
            apply_inline(p,text)
            i+=1;continue
        nm=re.match(r"^(\s*)\d+\.\s+(.*)",stripped)
        if nm:
            text=nm.group(2)
            p=doc.add_paragraph(style="List Number")
            p.paragraph_format.space_before=Pt(0)
            p.paragraph_format.space_after=Pt(2)
            apply_inline(p,text)
            i+=1;continue
        if stripped=="":
            i+=1;continue
        p=doc.add_paragraph()
        p.paragraph_format.space_before=Pt(0)
        p.paragraph_format.space_after=Pt(6)
        apply_inline(p,stripped)
        i+=1
    if table_buf: flush_table(doc,table_buf)
    if in_code: flush_code(doc,code_lines)
    doc.save(docx_path)
    print("[OK] Saved:",docx_path)


parse_and_build(
    "C:/Dev/flutterflowprojects/BiblioShare/docs/ARCHITECTURE.md",
    "C:/Dev/flutterflowprojects/BiblioShare/docs/ARCHITECTURE.docx"
)
